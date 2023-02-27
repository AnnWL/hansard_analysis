import pandas as pd
import numpy as np
import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json 
import datetime
import os, shutil
import os.path
import matplotlib.pyplot as plt
from matplotlib.dates import AutoDateLocator, AutoDateFormatter, date2num, DateFormatter
import matplotlib.dates as mdates
import textwrap
import time
import re
# Setting up paths
data_path = ('data/')

columns = ['idx', 'sitting_date', 'sitting_date_dt', 'asker_name', 'asker_party', 'asker_name_party','Is_PAP', 'asker_parliaments', 'askees','responder_name', 'responder_title', 'responder_name_title','title', 'link', 'ministry', 'question_comment', 'response_content', 'type','theme','question_topic_label']

def check_password():
    """Returns `True` if the user had the correct password."""

    def password_entered():
        """Checks whether a password entered by the user is correct."""
        if st.session_state["password"] == st.secrets["password"]:
            st.session_state["password_correct"] = True
            del st.session_state["password"]  # don't store password
        else:
            st.session_state["password_correct"] = False

    if "password_correct" not in st.session_state:
        # First run, show input for password.
        st.text_input(
            "Password", type="password", on_change=password_entered, key="password"
        )
        return False
    elif not st.session_state["password_correct"]:
        # Password not correct, show input + error.
        st.text_input(
            "Password", type="password", on_change=password_entered, key="password"
        )
        st.error("😕 Password incorrect")
        return False
    else:
        # Password correct.
        return True

def read_json(file_name): 
    with open(data_path+file_name) as json_file:
        output = json.load(json_file)
    return output

def save_json(json_obj, file_name): 
    with open(data_path+file_name, "w") as outfile:
        json.dump(json_obj, outfile)
    return 

@st.experimental_memo
def df_merge(overall_df_file, df_topic_file, party_acronyms_dict, columns): 
    df_all = pd.read_csv(data_path + overall_df_file)
    df_topic = pd.read_csv(data_path+df_topic_file)
    df = df_all.merge(df_topic[['idx','question_topic_label', 'theme']], on='idx', how = 'left')
    
    df.insert(3, 'asker_name_party', df.loc[:,'asker_name']+' ('+ df['asker_party'].\
               apply(lambda x:party_acronyms_dict[x])+')')
    
    df.insert(4,'Is_PAP', df.loc[:,'asker_party'].str.contains("People's Action Party"))
    df.fillna('',inplace = True)
    
    df['sitting_date_dt'] = pd.to_datetime(df['sitting_date'])
    
    df = df[df.sitting_date>='2018-01-01'].copy()
    
    df['responder_name_title'] = df['responder_name'] + ', ' + df['responder_title']
    
    return df[columns]

########
#Output#
########
mp_dict = read_json("mp_dict.json")

def get_impute_values(params_dict, origin_val):
    changed_params=dict()
    for key,val in params_dict.items():
        if key=='phrase':
            if val!='': 
                changed_params['phrase'] = val 
        elif val!= origin_val:
            changed_params[key] = val
    params_combi = set(changed_params.keys()) - {"time_ref_key","reference_date",'MP_name_party'}
    return params_combi, changed_params
    
def print_output(output_df, container):
    for i, row in output_df.iterrows(): 
        search_pos = row['link'].find('search/')+ 7
        new_link = row['link'][:search_pos] + '#/'+ row['link'][search_pos:]
        PQ_string = '**'+str(i)+'.** ['+row['question_comment']+']('+ new_link +')  \n'+\
        '**Sitting Date**: '+row['sitting_date']+'  \n'+'**Filed by MP**: '+row['asker_name_party']+'  \n'+\
        '**Responded by**: '+row['responder_name_title']
        container.markdown(PQ_string)
    return

def remove_special_chars(text):
    # Replace special characters with a space
    return re.sub(r'[^\w\s]', ' ', text)

@st.experimental_memo
def get_df_slice(df, params_dict):
    """Returns the sliced data according to users' impute parameters."""
    df_slice = df[df.sitting_date_dt > params_dict['reference_date']].copy()
    
    df_slice['question_comment_no_special_char'] = df_slice['question_comment'].apply(remove_special_chars)
    
    if params_dict['phrase'] != '': 
        df_slice = df_slice[df_slice.question_comment_no_special_char.str.lower().str.contains(' '+\
                           params_dict['phrase'].lower() +' ',regex = False)]
    if params_dict['agency'] != 'None':
        df_slice = df_slice[df_slice.ministry == params_dict['agency']]
    
    if params_dict['MP_name'] != 'None': 
        df_slice = df_slice[df_slice.asker_name == params_dict['MP_name'] ]
        
    if params_dict['theme'] != 'None': 
        df_slice = df_slice[df_slice['theme'] == params_dict['theme']]
        
    if params_dict['topic'] != 'None': 
        df_slice = df_slice[df_slice['question_topic_label'] == params_dict['topic']]

    df_slice = df_slice.sort_values(by = 'sitting_date', ascending = False).reset_index(drop= True)
    df_slice.index = df_slice.index+1
    
    return df_slice

def generate_summary_string(df_size, changed_params):
    """Returns a string comprising the number of PQs left after filtering through to users' impute parameters."""
    changed_params_keys = list(changed_params.keys())
    
    if df_size>0: 
        summary_str = "**"+str(df_size)+ " PQs "
    else: 
        summary_str = "**No PQs"
    
    if "phrase" in changed_params_keys:
        summary_str+=(" containing phrase/keyword '"+changed_params['phrase']+"'")
    
    summary_str += " filed"
    
    if "MP_name" in changed_params_keys: 
        summary_str+=" by "+ changed_params['MP_name'] 
        
    if "agency" in changed_params_keys: 
        summary_str+=" to "+ changed_params['agency'] 
    
    if "theme" in changed_params_keys: 
        summary_str+=" on theme '"+ changed_params['theme'] +"'"
        if "topic" in changed_params_keys:
            summary_str+= "and topic '" + changed_params['topic']+"'"
    else: 
        if "topic" in changed_params_keys:
            summary_str+= " on topic '" + changed_params['topic']+"'"
        
    summary_str+=" in the past " + changed_params['time_ref_key'] + ".**"
    return summary_str

def generate_time_series_title(changed_params):
    changed_params_keys = list(changed_params.keys())
    
    summary_str = 'Number of PQs'

    if "phrase" in changed_params_keys:
        summary_str+=(" containing phrase/keyword '"+changed_params['phrase']+"'")
    
    summary_str += " filed"
    
    if "MP_name" in changed_params_keys: 
        summary_str+=" by "+ changed_params['MP_name'] 
        
    if "agency" in changed_params_keys: 
        summary_str+=" to "+ changed_params['agency'] 
    
    if "theme" in changed_params_keys: 
        summary_str+=" on theme '"+ changed_params['theme'] +"'"
        if "topic" in changed_params_keys:
            summary_str+= "and topic '" + changed_params['topic']+"'"
    else: 
        if "topic" in changed_params_keys:
            summary_str+= "on topic '" + changed_params['topic']+"'"
        
    summary_str+=" in the past " + changed_params['time_ref_key'] 
    return summary_str

########
#Charts#
########


def freq_count(df_slice, parameter, title_val, top_n= 5): 
    data = pd.DataFrame(df_slice[parameter].value_counts().reset_index())
    data.columns =[parameter,'count']
    data['pct'] = round(data['count']*100/data['count'].sum()).astype(int)
    data.sort_values(by='count', ascending = False, inplace = True)
    data.index = data.index+1
    threshold = min(top_n, data.shape[0])
    return data.iloc[:threshold,:]
    
def bar_chart(df_slice, parameter, title_val, top_n= 5): 
    data = freq_count(df_slice, parameter, title_val)
    fig, ax = plt.subplots(figsize=(10, 4))
    if parameter=='asker_name':
        b1 = ax.barh(['\n'.join(textwrap.wrap(mp_dict[label],20)) for label in data[parameter]],data['count'].to_list(), color = 'tab:blue')
    else: 
        b1 = ax.barh(['\n'.join(textwrap.wrap(label,20)) for label in data[parameter]],data['count'].to_list(), color = 'tab:blue')
    ax.bar_label(b1, label_type='edge', labels=[" "+str(pct_val)+"%" for pct_val in data['pct']])
    ax.invert_yaxis()
    ax.set(ylabel='',xlabel='Number of PQs',title=title_val)
    return fig

def time_trend(df_slice, title_val, width_val = 31): 
    pq_counts = df_slice.resample(rule='M', on='sitting_date_dt').size().T.reset_index()
    pq_counts.columns = ['Sitting Date', 'Number of PQs']

    fig, ax = plt.subplots(figsize=(10, 4))
    ax.bar(pq_counts['Sitting Date'], pq_counts['Number of PQs'], color = 'tab:blue',width = width_val)
    ax.xaxis.set_major_formatter(mdates.ConciseDateFormatter(ax.xaxis.get_major_locator()))

    months = mdates.MonthLocator((1,2,3,4,5,6,7,8,9,10,11,12))
    ax.xaxis.set_minor_locator(months)
    ax.set(xlabel='Sitting Date',ylabel='Number of PQs',\
           title=title_val)
    return fig

def time_trend_PAP(df_slice, title_val, width_val = 31): 
    t_index = pd.DatetimeIndex(pd.date_range(start=df_slice.sitting_date_dt.min(), end=df_slice.sitting_date_dt.max()))
    pq_counts_pap = df_slice[df_slice.Is_PAP].resample(rule='M',\
                                                       on='sitting_date_dt').size().reindex(t_index).fillna(0).reset_index()
    pq_counts_pap.columns = ['Sitting Date', 'Num_PQs']
    pq_counts_pap_not = df_slice[~df_slice.Is_PAP].resample(rule='M',\
                                                   on='sitting_date_dt').size().reindex(t_index).fillna(0).reset_index()
    pq_counts_pap_not.columns = ['Sitting Date', 'Num_PQs']

    fig, ax = plt.subplots(figsize=(10, 2))
    p1 = ax.bar(pq_counts_pap_not['Sitting Date'], pq_counts_pap_not['Num_PQs'], color = 'tab:orange',width = width_val)
    p2 = ax.bar(pq_counts_pap['Sitting Date'], pq_counts_pap['Num_PQs'], bottom = pq_counts_pap_not['Num_PQs'], \
                color = 'tab:blue',width = width_val)
    ax.legend((p1[0], p2[0]), ('Non-PAP MPs', 'PAP MPs'))
    ax.xaxis.set_major_formatter(mdates.ConciseDateFormatter(ax.xaxis.get_major_locator()))
    months = mdates.MonthLocator((1,2,3,4,5,6,7,8,9,10,11,12))
    ax.xaxis.set_minor_locator(months)
    ax.set(xlabel='Sitting Date',ylabel='Number of PQs',\
           title=title_val)
    return fig

#######
#Brief#
#######

def make_archive(source, destination):
    base = os.path.basename(destination)
    name = base.split('.')[0]
    format = base.split('.')[1]
    archive_from = os.path.dirname(source)
    archive_to = os.path.basename(source.strip(os.sep))
    shutil.make_archive(name, format, archive_from, archive_to)
    shutil.move('%s.%s'%(name,format), destination)
    return
        
def generate_pq_docx(i,row): 
    doc = Document()
    p = doc.add_paragraph('#'+str(i+1)+'\n')
    p.add_run('Sitting Date: ').bold = True
    p.add_run(row['sitting_date'] + '\n')

    p.add_run('Asker:  ').bold = True
    p.add_run(row['asker_name_party'] +'\n')

    p.add_run('Responder: ').bold = True
    p.add_run(row['responder_name'] + ', ' + row['responder_title']+'\n')
    p.add_run('Title: ').bold = True
    p.add_run(row['title']+'\n')

    if row['question_topic_label']!='':
        p.add_run('Topic: ').bold = True
        p.add_run(row['question_topic_label'])
    else: 
        p.add_run('Topic: ').bold = True
        p.add_run('Not available')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run('Question: ').bold = True
    p.add_run(row['question_comment'])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run('Response: ').bold = True
    p.add_run(row['response_content'])

    doc.add_paragraph("_____________________________________________")
    p2 = doc.add_paragraph('Link: '+ str(row['link']))
    return doc

def generate_folder(output_df, changed_params): 
    folder_name_val ='PQ_Background_Briefs_for'
    
    for key, value in changed_params.items(): 
        if key not in ['reference_date','MP_name']:
            folder_name_val+='_'+value
    
    current_directory = os.getcwd()
    output_directory = os.path.join(current_directory, folder_name_val)

    if not os.path.exists(output_directory):
        os.makedirs(output_directory)
    
    for i, row in output_df.iterrows(): 
        doc = generate_pq_docx(i,row)
        title = 'PQ_'+str(i)+'_'+str(row['sitting_date'])+'.docx'
        doc.save(folder_name_val+'/'+title)
    
    make_archive(output_directory, output_directory+'.zip')
    
    #output_container = st.sidebar.container()
    #output_container.write('**Output**')
    
    #with open(output_directory+'.zip', "rb") as fp:
    #    btn = output_container.download_button(
    #        label="Download Background Briefs",
    #        data=fp,
    #        file_name=folder_name_val+".zip",
    #        mime="application/zip")
        
       
    return output_directory
